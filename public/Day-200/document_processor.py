import os
import logging
from typing import Optional
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handle document processing and text extraction."""
    
    def __init__(self):
        """Initialize the document processor."""
        pass
    
    def extract_text(self, filepath: str) -> Optional[str]:
        """
        Extract text from various document formats.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            file_extension = os.path.splitext(filepath)[1].lower()
            
            if file_extension == '.txt':
                return self._extract_from_txt(filepath)
            elif file_extension == '.pdf':
                if PYMUPDF_AVAILABLE:
                    return self._extract_from_pdf(filepath)
                else:
                    logger.error("PDF processing not available - PyMuPDF not installed")
                    return None
            elif file_extension == '.docx':
                if DOCX_AVAILABLE:
                    return self._extract_from_docx(filepath)
                else:
                    logger.error("DOCX processing not available - python-docx not installed")
                    return None
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return None
    
    def _extract_from_txt(self, filepath: str) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as file:
                        text = file.read()
                        if text.strip():
                            return text
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file with any encoding: {filepath}")
            return None
            
        except Exception as e:
            logger.error(f"Error reading TXT file {filepath}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, filepath: str) -> Optional[str]:
        """Extract text from PDF file using PyMuPDF."""
        try:
            text = ""
            doc = fitz.open(filepath)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += page_text + "\n"
            
            doc.close()
            
            if text.strip():
                return text.strip()
            else:
                logger.warning(f"No text extracted from PDF: {filepath}")
                return None
                
        except Exception as e:
            logger.error(f"Error reading PDF file {filepath}: {str(e)}")
            return None
    
    def _extract_from_docx(self, filepath: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(filepath)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                return text.strip()
            else:
                logger.warning(f"No text extracted from DOCX: {filepath}")
                return None
                
        except Exception as e:
            logger.error(f"Error reading DOCX file {filepath}: {str(e)}")
            return None
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        # Join lines with single spaces, but preserve paragraph breaks
        cleaned_text = ""
        for i, line in enumerate(cleaned_lines):
            if i > 0:
                # Add paragraph break for short lines (likely headers/breaks)
                if len(cleaned_lines[i-1]) < 80:
                    cleaned_text += "\n\n"
                else:
                    cleaned_text += " "
            cleaned_text += line
        
        return cleaned_text.strip()
